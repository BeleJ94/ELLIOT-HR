from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "docs"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 108, 128)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
INK = RGBColor(31, 41, 55)
GREEN = RGBColor(22, 101, 52)
GOLD = RGBColor(122, 90, 0)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.text = ""
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_borders(table, color="D7DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_document(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in [
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 12.5, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("ELLIOT-HR | " + title)
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Document professionnel genere pour ELLIOT-HR")
    set_run_font(r, size=8, color=MUTED)


def add_cover(doc, kicker, title, subtitle, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.upper())
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    table = doc.add_table(rows=len(meta), cols=2)
    set_table_widths(table, [1.8, 4.2])
    set_borders(table, "DDE5EE", "4")
    for i, (label, value) in enumerate(meta):
        shade_cell(table.cell(i, 0), LIGHT_BLUE)
        shade_cell(table.cell(i, 1), "FFFFFF")
        set_cell_text(table.cell(i, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(i, 1), value)
    doc.add_page_break()


def add_callout(doc, title, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    set_borders(table, "D7DEE8", "6")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=INK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    set_borders(table)
    for idx, header in enumerate(headers):
        shade_cell(table.cell(0, idx), LIGHT_BLUE)
        set_cell_text(table.cell(0, idx), header, bold=True, color=DARK_BLUE, size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    return table


def p(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def h1(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def presentation():
    doc = Document()
    style_document(doc, "Presentation du logiciel")
    add_cover(
        doc,
        "Presentation du logiciel",
        "ELLIOT-HR",
        "Suite professionnelle de gestion du capital humain, de la paie et des operations RH",
        [
            ("Document", "Presentation du logiciel"),
            ("Public cible", "Direction generale, Direction RH, managers, responsables paie"),
            ("Version", "Pack fonctionnel 2026"),
            ("Positionnement", "ERP RH modulaire pour entreprises multi-sites"),
        ],
    )

    h1(doc, "1. Vision generale")
    p(doc, "ELLIOT-HR est une suite de gestion RH concue pour centraliser les donnees collaborateurs, structurer les processus administratifs et donner aux responsables une vision claire des operations humaines. Le logiciel couvre le cycle de vie du collaborateur, du dossier employe a la paie, en passant par les contrats, les presences, les conges, les formations et les declarations.")
    add_callout(doc, "Promesse produit", "Transformer les operations RH quotidiennes en processus controles, tracables et exploitables par la direction.")

    h1(doc, "2. Proposition de valeur")
    add_table(
        doc,
        ["Enjeu", "Reponse ELLIOT-HR", "Impact attendu"],
        [
            ("Dossiers RH disperses", "Centralisation des fiches employes, documents et affectations", "Gain de temps et meilleure qualite des donnees"),
            ("Suivi operationnel insuffisant", "Dashboard, KPI, alertes et details exportables", "Pilotage plus rapide et decisions documentees"),
            ("Paie sensible", "Parametrage paie, simulation, bulletins et journal", "Reduction des erreurs et meilleure anticipation"),
            ("Presences et conges complexes", "Pointage journalier, validation, historique et rapports", "Controle terrain et transparence"),
            ("Competences non suivies", "Catalogue, sessions multi-jours, participants et historiques formation", "Developpement RH mesurable"),
        ],
        [1.55, 2.55, 2.4],
    )

    h1(doc, "3. Modules fonctionnels")
    modules = [
        ("Tableau de bord", "Vue globale des KPI, notifications, graphiques, details en modales, exports PDF/Excel."),
        ("Collaborateurs", "Gestion des employes, affectations, contrats, pieces et historique individuel."),
        ("Temps & absences", "Presences, pointage, controle des journees, conges, approbations et rapports."),
        ("Developpement RH", "Catalogue formation, sessions multi-jours, participants, appels journaliers et certificats."),
        ("Paie & obligations", "Periodes de paie, calcul, simulation net vers base, bulletins, declarations et exports."),
        ("Organisation", "Entreprises, departements, postes, sites/agences et referentiels."),
        ("Administration", "Utilisateurs, roles, droits d'acces et securite applicative."),
    ]
    add_table(doc, ["Module", "Role dans le logiciel"], modules, [1.7, 4.8])

    h1(doc, "4. Parcours utilisateur cible")
    h2(doc, "Direction et responsables")
    add_bullets(doc, [
        "Consulter les indicateurs cles sur le tableau de bord.",
        "Ouvrir les details des KPI dans des modales lisibles.",
        "Exporter les analyses en PDF ou Excel pour comite, audit ou suivi interne.",
    ])
    h2(doc, "Equipe RH")
    add_bullets(doc, [
        "Creer et tenir a jour les dossiers employes.",
        "Suivre les contrats, absences, formations et documents.",
        "Produire les rapports operationnels et historiques individuels.",
    ])
    h2(doc, "Paie et obligations")
    add_bullets(doc, [
        "Preparer les periodes, calculer les bulletins et controler les retenues.",
        "Simuler un salaire a partir d'un net cible.",
        "Generer declarations et justificatifs exportables.",
    ])

    h1(doc, "5. Points forts de l'interface")
    add_bullets(doc, [
        "Organisation du menu par parcours metier: collaborateurs, temps, formation, paie, organisation.",
        "Rendu ERP professionnel, dense mais lisible, adapte a une utilisation quotidienne.",
        "Actions rapides, filtres, tableaux et modales de detail.",
        "Exports PDF/Excel integres aux workflows critiques.",
        "Historique et tracabilite par employe et par module.",
    ])

    h1(doc, "6. Securite et gouvernance")
    p(doc, "Le logiciel applique une logique d'acces par role. Les utilisateurs super-admin, admin RH, manager et employe disposent de vues differentes selon leur niveau d'autorisation. Les actions sensibles sont protegees par jeton CSRF et journalisees lorsque le flux le prevoit.")
    add_table(
        doc,
        ["Role", "Acces principal"],
        [
            ("Super admin", "Pilotage global, entreprises, configuration et administration plateforme."),
            ("Admin RH", "Gestion RH complete dans le perimetre de l'entreprise."),
            ("Manager", "Suivi equipe, presences, conges, formations et dossiers autorises."),
            ("Employe", "Acces limite aux fonctions personnelles prevues."),
        ],
        [1.7, 4.8],
    )

    h1(doc, "7. Conclusion")
    p(doc, "ELLIOT-HR se positionne comme une plateforme RH operationnelle et evolutive. Sa force reside dans l'association entre donnees RH centralisees, processus quotidiens controles, indicateurs exploitables et exports professionnels. Le logiciel constitue une base solide pour digitaliser les pratiques RH d'une organisation moderne.")

    path = OUT / "ELLIOT-HR_Presentation_du_logiciel.docx"
    doc.save(path)
    return path


def guide():
    doc = Document()
    style_document(doc, "Guide utilisateur complet")
    add_cover(
        doc,
        "Guide utilisateur complet",
        "ELLIOT-HR",
        "Manuel operationnel pour administrateurs RH, managers, paie et utilisateurs autorises",
        [
            ("Document", "Guide utilisateur complet"),
            ("Public cible", "Super admin, admin RH, manager, employe"),
            ("Usage", "Formation interne, support utilisateur, reference operationnelle"),
            ("Version", "Edition 2026"),
        ],
    )

    h1(doc, "1. Introduction")
    p(doc, "Ce guide explique comment utiliser ELLIOT-HR au quotidien. Il decrit les modules, les parcours principaux, les bonnes pratiques et les controles a effectuer pour garantir une gestion RH fiable.")
    add_callout(doc, "Principe de navigation", "Le menu principal est organise par parcours: Tableau de bord, Collaborateurs, Temps & absences, Developpement RH, Paie & obligations, Organisation et Administration.")

    h1(doc, "2. Connexion et environnement")
    add_steps(doc, [
        "Ouvrir l'adresse du logiciel dans le navigateur.",
        "Saisir l'adresse email professionnelle et le mot de passe.",
        "Verifier le role affiche dans la barre superieure.",
        "Utiliser la recherche rapide pour ouvrir un module.",
    ])
    h2(doc, "Bonnes pratiques de securite")
    add_bullets(doc, [
        "Ne pas partager son compte utilisateur.",
        "Se deconnecter apres usage sur un poste partage.",
        "Signaler immediatement tout acces ou comportement inhabituel.",
    ])

    h1(doc, "3. Tableau de bord")
    p(doc, "Le tableau de bord fournit une vue executive des effectifs, contrats, conges, presences, notifications, abonnements et autres indicateurs utiles selon le role.")
    add_steps(doc, [
        "Consulter les KPI principaux.",
        "Cliquer sur un KPI, une notification ou un graphique pour afficher les details.",
        "Utiliser les boutons PDF ou Excel dans la modale pour extraire le rapport.",
        "Traiter les priorites RH signalees par les alertes.",
    ])

    h1(doc, "4. Collaborateurs")
    h2(doc, "Employes")
    add_steps(doc, [
        "Ouvrir Collaborateurs > Employes.",
        "Filtrer par entreprise, departement, poste ou statut.",
        "Cliquer sur un employe pour ouvrir son dossier.",
        "Verifier les informations personnelles, l'affectation, les documents et l'historique formation.",
        "Utiliser Modifier pour mettre a jour le dossier.",
    ])
    h2(doc, "Contrats")
    add_steps(doc, [
        "Ouvrir Collaborateurs > Contrats.",
        "Creer ou consulter un contrat.",
        "Suivre les contrats actifs, expires ou a risque.",
        "Generer les documents et archiver les pieces signees lorsque prevu.",
    ])

    h1(doc, "5. Temps & absences")
    h2(doc, "Presences")
    add_steps(doc, [
        "Selectionner la date de pointage.",
        "Filtrer les employes si necessaire.",
        "Renseigner statut, heure d'arrivee, heure de depart et note.",
        "Enregistrer la journee.",
        "Cloturer ou verrouiller selon les droits disponibles.",
    ])
    h2(doc, "Conges")
    add_steps(doc, [
        "Ouvrir Temps & absences > Conges.",
        "Creer une demande ou consulter l'historique.",
        "Pour les managers et RH, ouvrir la validation.",
        "Valider ou refuser avec motif selon le workflow.",
        "Suivre les soldes et le calendrier des absences.",
    ])

    h1(doc, "6. Developpement RH: formations")
    p(doc, "Le module Formations sert a suivre les mises a niveau, formations internes, formations externes et sessions multi-jours.")
    h2(doc, "Catalogue")
    add_steps(doc, [
        "Ouvrir Developpement RH > Formations.",
        "Cliquer sur Catalogue.",
        "Renseigner l'intitule, le code, le domaine, la duree et les objectifs.",
        "Enregistrer la formation.",
    ])
    h2(doc, "Session multi-jours")
    add_steps(doc, [
        "Cliquer sur Nouvelle session.",
        "Choisir la formation du catalogue.",
        "Indiquer titre, date debut, date fin, horaires, lieu et formateur.",
        "Definir le seuil de validation, par exemple 80%.",
        "Enregistrer: les journees sont generees automatiquement.",
    ])
    h2(doc, "Participants et presence")
    add_steps(doc, [
        "Ouvrir la fiche session.",
        "Ajouter les employes invites.",
        "Selectionner chaque jour de formation.",
        "Faire l'appel: present, retard, absent ou excuse.",
        "Enregistrer l'appel de la journee.",
        "Finaliser la session pour calculer le statut final et le certificat.",
    ])
    add_table(
        doc,
        ["Statut", "Signification"],
        [
            ("Invite", "Employe ajoute a la session mais non encore finalise."),
            ("Valide", "Taux de presence egal ou superieur au seuil."),
            ("Non valide", "Participation insuffisante selon le seuil defini."),
            ("Absent", "Aucune presence effective enregistree."),
            ("Excuse", "Absence justifiee selon le suivi RH."),
        ],
        [1.6, 4.9],
    )

    h1(doc, "7. Paie & obligations")
    h2(doc, "Paie")
    add_steps(doc, [
        "Creer ou ouvrir une periode de paie.",
        "Verifier les employes et les elements de paie.",
        "Lancer le calcul.",
        "Controler le journal de paie.",
        "Fermer la periode lorsque les donnees sont validees.",
        "Generer les bulletins PDF ou exporter le journal.",
    ])
    h2(doc, "Simulation de paie")
    add_steps(doc, [
        "Ouvrir le module de simulation.",
        "Saisir le net cible.",
        "Lancer le calcul.",
        "Analyser le salaire de base, le brut, les retenues, le net et le cout employeur.",
        "Exporter la simulation en PDF si necessaire.",
    ])
    h2(doc, "Declarations")
    add_steps(doc, [
        "Generer les declarations a partir des periodes de paie.",
        "Verifier les montants dus.",
        "Mettre a jour le statut de paiement.",
        "Joindre une preuve de paiement lorsque disponible.",
        "Exporter le dossier en PDF ou Excel.",
    ])

    h1(doc, "8. Organisation")
    add_bullets(doc, [
        "Entreprises: gestion du profil entreprise, abonnements et informations legales.",
        "Departements: structure interne et rattachements.",
        "Postes: referentiel des fonctions et affectations.",
    ])

    h1(doc, "9. Administration")
    h2(doc, "Utilisateurs")
    add_steps(doc, [
        "Creer les comptes utilisateurs.",
        "Associer le role approprie.",
        "Lier un compte a un employe si necessaire.",
        "Activer, bloquer ou reinitialiser un mot de passe selon les droits.",
    ])
    add_table(
        doc,
        ["Role", "Responsabilites usuelles"],
        [
            ("Super admin", "Administration globale de la plateforme et des entreprises."),
            ("Admin RH", "Gestion RH complete du perimetre entreprise."),
            ("Manager", "Suivi equipe, presences, conges et formations."),
            ("Employe", "Acces limite aux informations et actions personnelles."),
        ],
        [1.5, 5.0],
    )

    h1(doc, "10. Exports et rapports")
    p(doc, "Les exports permettent de produire des justificatifs, des rapports de pilotage ou des fichiers de travail. Les formats disponibles dependent du module: PDF pour la presentation, Excel pour l'analyse et les listes.")
    add_bullets(doc, [
        "PDF: privilegier pour transmission, impression ou archivage.",
        "Excel: privilegier pour retraitement, filtrage ou controle.",
        "Verifier les filtres avant extraction pour eviter les rapports incomplets.",
    ])

    h1(doc, "11. Controles qualite recommandes")
    add_table(
        doc,
        ["Frequence", "Controle"],
        [
            ("Chaque jour", "Verifier les presences et alertes du tableau de bord."),
            ("Chaque semaine", "Controler conges en attente, contrats a risque et formations en cours."),
            ("Chaque mois", "Verifier paie, declarations, documents et exports de synthese."),
            ("Chaque trimestre", "Analyser formations realisees, competences et historiques employes."),
        ],
        [1.5, 5.0],
    )

    h1(doc, "12. Depannage rapide")
    add_table(
        doc,
        ["Probleme", "Action conseillee"],
        [
            ("Un bouton ne repond pas", "Actualiser la page et verifier la connexion."),
            ("Une modale semble mal affichee", "Faire un rafraichissement dur pour recharger CSS/JS."),
            ("Un employe n'apparait pas", "Verifier son statut, son entreprise et les filtres actifs."),
            ("Un export est vide", "Verifier le perimetre, les dates et les donnees sources."),
            ("Acces refuse", "Verifier le role utilisateur et le rattachement entreprise."),
        ],
        [2.0, 4.5],
    )

    h1(doc, "13. Glossaire")
    add_table(
        doc,
        ["Terme", "Definition"],
        [
            ("Perimetre", "Ensemble des donnees accessibles selon le role et l'entreprise."),
            ("Session", "Occurrence planifiee d'une formation, avec dates et participants."),
            ("Seuil de validation", "Taux de presence minimum requis pour valider une formation."),
            ("Journal de paie", "Liste consolidee des bulletins et montants d'une periode."),
            ("Declaration", "Document fiscal ou social genere a partir de la paie."),
        ],
        [1.8, 4.7],
    )

    path = OUT / "ELLIOT-HR_Guide_utilisateur_complet.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    print(presentation())
    print(guide())
